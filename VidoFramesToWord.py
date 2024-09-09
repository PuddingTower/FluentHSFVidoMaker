import cv2
import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

def set_font_style(paragraph, font_name="Times New Roman", font_size=Pt(12)):
    """
    设置段落中的字体样式
    """
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        # 设置东亚字体（用于中文字体的兼容性）
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)

def extract_frames_to_word(video_path, output_docx_path, num_frames=10):
    # 打开视频文件
    video = cv2.VideoCapture(video_path)
    
    # 检查视频是否成功打开
    if not video.isOpened():
        print(f"Failed to open video file: {video_path}")
        return

    # 获取视频的总帧数和帧率
    total_frames = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = video.get(cv2.CAP_PROP_FPS)

    # 检查帧率和总帧数是否有效
    if fps == 0 or total_frames == 0:
        print(f"Invalid video file: {video_path}")
        return

    duration = total_frames / fps  # 视频的总时长（秒）

    # 确保提取的帧数不超过视频的总帧数
    if num_frames > total_frames:
        num_frames = total_frames

    # 计算每个时间点对应的帧数
    frame_interval = total_frames // num_frames
    
    # 创建Word文档
    doc = Document()
    doc.add_heading('Extracted Frames from Video', 0)
    
    # 提取每个时间点的帧并插入到Word文档中
    for i in range(num_frames):
        frame_number = i * frame_interval
        video.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
        ret, frame = video.read()
        if ret:
            # 保存帧为图片文件
            frame_time = frame_number / fps
            image_filename = f"frame_{i+1}.jpg"
            cv2.imwrite(image_filename, frame)
            
            # 将时间和图片插入Word文档
            paragraph = doc.add_paragraph(f"{frame_time:.2f}s:")
            set_font_style(paragraph, font_name="Times New Roman", font_size=Pt(12))
            doc.add_picture(image_filename, width=Inches(2))  # 调整图片宽度
            
            # 删除临时图片文件
            os.remove(image_filename)
    
    # 保存Word文档
    doc.save(output_docx_path)
    video.release()
    print(f"Word document saved as {output_docx_path}")

if __name__ == "__main__":
    # 确定程序运行目录
    if getattr(sys, 'frozen', False):
        # 如果是打包后的可执行文件，使用 sys.executable 获取实际目录
        current_dir = os.path.dirname(sys.executable)
    else:
        # 否则使用 __file__ 获取脚本文件的目录
        current_dir = os.path.dirname(os.path.abspath(__file__))

    # 在当前目录下查找 output.mp4
    video_file = os.path.join(current_dir, "output.mp4")
    output_docx = os.path.join(current_dir, "output_frames.docx")

    if not os.path.exists(video_file):
        print(f"Video file {video_file} does not exist.")
    else:
        extract_frames_to_word(video_file, output_docx)
